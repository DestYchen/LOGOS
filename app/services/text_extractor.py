from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

_IMAGE_SUFFIXES = {".pdf", ".png", ".jpg", ".jpeg"}
_IMAGE_MIME_TYPES = {"application/pdf", "image/png", "image/jpeg"}


@dataclass(frozen=True)
class TextExtractionResult:
    """Container for parser based text extraction results."""

    text: str
    parser: str


def requires_ocr(path: Path, mime: Optional[str] = None) -> bool:
    """Return True when OCR should be executed for the provided file."""

    suffix = path.suffix.lower()
    mime_type = (mime or "").split(";", 1)[0].strip().lower()
    return suffix in _IMAGE_SUFFIXES or mime_type in _IMAGE_MIME_TYPES


def extract_text(path: Path, mime: Optional[str] = None) -> Optional[TextExtractionResult]:
    """Return textual content for non image documents when possible."""

    if requires_ocr(path, mime):
        return None

    suffix = path.suffix.lower()
    mime_type = (mime or "").split(";", 1)[0].strip().lower()

    try:
        if _is_docx(suffix, mime_type):
            text = _extract_docx(path)
            parser = "docx"
        elif _is_xlsx(suffix, mime_type):
            text = _extract_xlsx(path)
            parser = "xlsx"
        elif _is_plain_text(suffix, mime_type):
            text = _extract_plain_text(path)
            parser = "text"
        else:
            return None
    except Exception:  # pragma: no cover - best effort fall back
        logger.exception("Failed to extract text from %s", path)
        return None

    cleaned = text.strip()
    if not cleaned:
        return None
    return TextExtractionResult(text=cleaned, parser=parser)


def _is_docx(suffix: str, mime_type: str) -> bool:
    return suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _is_xlsx(suffix: str, mime_type: str) -> bool:
    return suffix in {".xlsx", ".xlsm"} or mime_type in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    }


def _is_plain_text(suffix: str, mime_type: str) -> bool:
    if mime_type.startswith("text/"):
        return True
    return suffix in {".txt", ".csv", ".tsv", ".json", ".md"}


def _extract_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    from docx import Document  # type: ignore import-not-found

    document = Document(path)
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append("\t".join(cells))
    return "\n".join(chunks)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook  # type: ignore import-not-found

    workbook = load_workbook(filename=path, data_only=True, read_only=True)
    try:
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell not in (None, "")]
                if values:
                    chunks.append("\t".join(values))
        return "\n".join(chunks)
    finally:
        workbook.close()
